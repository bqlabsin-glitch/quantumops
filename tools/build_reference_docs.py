from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
NAVY = "102A43"
BLUE = "2563EB"
CYAN = "0891B2"
INK = "172B4D"
MUTED = "5B677A"
LIGHT = "EAF2FF"
PALE = "F4F7FB"
WHITE = "FFFFFF"


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def shade_paragraph(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = pPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        pPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = "Aptos"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Aptos")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Aptos")
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    if color is not None: run.font.color.rgb = RGBColor.from_string(color)


def setup(doc, running_label):
    s = doc.sections[0]
    s.page_width, s.page_height = Inches(8.5), Inches(11)
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1)
    s.header_distance = s.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name, normal.font.size = "Aptos", Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, before, after, color in (
        ("Title", 27, 0, 8, NAVY), ("Subtitle", 13, 0, 16, MUTED),
        ("Heading 1", 17, 16, 8, NAVY), ("Heading 2", 13, 12, 6, BLUE),
        ("Heading 3", 11.5, 8, 4, CYAN)):
        st = styles[name]
        st.font.name, st.font.size = "Aptos Display", Pt(size)
        st.font.bold = name not in ("Subtitle",)
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        st = styles[name]
        st.font.name, st.font.size = "Aptos", Pt(10.5)
        st.paragraph_format.left_indent = Inches(.5)
        st.paragraph_format.first_line_indent = Inches(-.25)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.10
    hp = s.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(hp.add_run(running_label.upper()), 8.5, True, MUTED)
    fp = s.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(fp.add_run("BQ Labs | Internal planning reference | 30 August 2026"), 8, False, MUTED)


def title(doc, kicker, name, subtitle, status):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    font(p.add_run(kicker.upper()), 9, True, CYAN)
    p = doc.add_paragraph(style="Title")
    p.add_run(name)
    p = doc.add_paragraph(style="Subtitle")
    p.add_run(subtitle)
    for label, value in (("Status", status), ("Purpose", "Living reference for design, product, and implementation decisions")):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        shade_paragraph(p, PALE)
        font(p.add_run(label + ": "), 9.5, True, NAVY)
        font(p.add_run(value), 9.5, False, INK)
    doc.add_paragraph()


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.add_run(item)


def add_steps(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    set_repeat_table_header(t.rows[0])
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        shade(cell, NAVY)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        font(p.add_run(h), 9, True, WHITE)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell_margins(cells[i])
            if ri % 2: shade(cells[i], PALE)
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            font(p.add_run(str(value)), 9, False, INK)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths): row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def callout(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(.12)
    p.paragraph_format.right_indent = Inches(.12)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    shade_paragraph(p, LIGHT)
    font(p.add_run(label + ": "), 10, True, NAVY)
    font(p.add_run(text), 10, False, INK)


def add_final_design_blueprint(doc, compact=False):
    doc.add_heading("Approved experience direction - 5 September 2026", 1)
    doc.add_paragraph("BQ Labs will use an Apple-inspired but independently designed experience system. The standard is professional restraint: clarity, hierarchy, original BQ Labs assets, excellent responsive behaviour, and motion that supports rather than competes with the user's work.")
    add_table(doc, ["Foundation", "Approved direction"], [
        ("Themes", "Light-dominant public site with selected cinematic dark sections; Quantum OPS light by default with complete dark-theme support."),
        ("Typography", "Self-hosted Manrope for expressive display text and Inter for interfaces, tables, forms, and body copy; system fallbacks."),
        ("Colour", "Near-black and soft-white neutrals, electric-blue primary accent, restrained cyan-violet gradients, semantic green/amber/red only."),
        ("Surfaces", "12-16px radii, thin neutral borders, subtle shadows, limited translucency, flatter high-density tables and forms."),
        ("Navigation", "Slim sticky public header; role-aware application sidebar/top bar; collapsible tablet shell; mobile bottom navigation."),
        ("Dashboards", "Role-based calm command centre: concise summary, urgent actions, few meaningful metrics, operational lists/tables, charts only for useful trends."),
        ("Motion", "Subtle transitions and feedback, sparing public scroll reveals, no workflow delays, no decorative table/form animation, reduced-motion support."),
    ], [1.35, 5.15])
    if compact:
        add_bullets(doc, ["The shared system applies to BQ Labs, Quantum OPS, and future products through reusable design tokens and accessible components.", "Product-specific accents are allowed within controlled accessibility and brand rules.", "The user has delegated remaining detailed design choices to the implementation team and will review the functioning application for adjustments."])
        return
    doc.add_heading("BQ Labs public-page architecture", 2)
    add_steps(doc, ["Minimal hero with one strong BQ Labs promise.", "Large Quantum OPS showcase using original product imagery.", "Short workflow or outcome story.", "Preview of future BQ Labs products.", "Security, privacy, and trust narrative without public diagnostics.", "Focused final call to action.", "Restrained company, support, privacy, and legal footer."])
    doc.add_heading("Quantum OPS interaction standards", 2)
    add_bullets(doc, ["Information density adapts to role and task: spacious onboarding and summaries, compact but readable operational tables.", "Every screen has one dominant next action, clear empty/loading/error states, and keyboard/touch parity.", "Charts require a decision-oriented purpose and accessible text/table equivalents.", "Mobile supports the complete critical path: assignments, status, blockers, UAT, approvals, leave, and client actions.", "Avoid copied Apple assets, layout replication, gratuitous glass, glow, novelty cursors, autoplay media, and motion-dependent comprehension.", "Shared tokens cover type, colour, spacing, radii, elevation, motion, breakpoints, density, focus, and semantic states. Future products extend rather than fork them."])
    doc.add_heading("Design verification gates", 2)
    add_bullets(doc, ["Responsive review at 360px, 768px, 1024px, and wide desktop layouts.", "WCAG-conscious contrast, visible focus, semantic landmarks, keyboard flow, labelled controls, accessible errors, and reduced motion.", "Performance budgets for fonts, images, JavaScript, layout stability, and interaction latency.", "Role-based usability walkthroughs for team member, Tester, Team Lead, client, senior management, Organization Admin, and BQ Labs Platform Admin.", "Original visual assets and copy reviewed for BQ Labs ownership, professional tone, and factual claims."])


def add_consolidated_phase2_decisions(doc):
    doc.add_heading("Consolidated Phase 2 decisions - 5 September 2026", 1)
    doc.add_heading("Tester and UAT model", 2)
    add_bullets(doc, ["The working team member selects Testers for the task.", "Every Tester must accept before raising observations; response is due by the next project working day.", "A task may have multiple Testers but exactly one Main Tester.", "Observation lifecycle: Open -> Fix In Progress -> Ready for Retest -> Retesting -> Resolved, with Rejected and Duplicate terminal alternatives.", "Any accepted Tester assigned to the task may retest and resolve an observation.", "Severity levels are Critical, High, Medium, and Low; response targets are project-configurable.", "A task cannot become Completed while any UAT observation remains open; no override is permitted."])
    doc.add_heading("Blockers, tasks, and ownership", 2)
    add_bullets(doc, ["Assignee, accepted Tester, Task Owner, and Team Lead may create blockers. Blocker owner proposes resolution; assignee or Task Owner confirms closure.", "Tasks cannot be permanently deleted through the normal application; Task Owner may archive them.", "All internal project members may create tasks, and the creator becomes Task Owner. Clients cannot create tasks directly.", "Task Owner controls scope, description, acceptance criteria, review requirement, priority, dates, and assignment; assignee controls execution status, blockers, progress, and ETA requests; Team Lead may change planning fields with audit.", "Support blocking dependencies and related-task links. Starting with an unresolved blocking dependency produces a warning but is allowed."])
    doc.add_heading("Files, comments, and client boundaries", 2)
    add_bullets(doc, ["Tasks and UAT observations support private attachments with project limits, type/size validation, malware scanning, client visibility, expiring downloads, and audit.", "Internal and Client-visible comments are separate immutable-audience channels; correcting audience requires an audited replacement.", "Clients may view permitted tasks/fields, use Client-visible comments, approve configured project leave impact, review only when task-authorized, participate in permitted UAT verification, and download client-visible files.", "Clients cannot alter internal status, assignment, planning, roles, or permissions."])
    doc.add_heading("Reporting and platform behaviour", 2)
    add_bullets(doc, ["Client progress uses completed-task weighting based on effort or complexity; in-progress work does not inflate completion, and milestone health is separate.", "Senior management sees authorized portfolio health, ETA risk, blockers, capacity, leave impact, UAT ageing, assignment delays, delivery trends, and intervention needs.", "Phase 2 notifications are in-app plus essential email. SMS, WhatsApp, and mobile push are deferred.", "Dashboards use efficient API refresh and near-real-time notification badges rather than fully streamed data.", "Search is permission-filtered before results or suggestions are returned; client search remains within assigned projects and visible fields."])
    doc.add_heading("Security, export, retention, and compatibility", 2)
    add_bullets(doc, ["Authentication includes verified email, secure reset, session/device management, rate limits, lockout protection, optional TOTP MFA, and re-authentication for sensitive actions; enterprise SSO is deferred.", "Authorized CSV/Excel exports enforce row and field permissions, size limits, expiring downloads, and audit; clients receive only authorized project data.", "Operational and audit records are retained by default. Permanent deletion is restricted to BQ Labs administration for approved legal/privacy cases.", "Support current and previous major Chrome, Edge, Firefox, and Safari versions, including mobile Chrome and Safari; Internet Explorer is excluded."])
    doc.add_heading("Release sequence", 2)
    add_steps(doc, ["BQ Labs internal pilot.", "One controlled client pilot.", "Security, performance, accessibility, and usability correction cycle.", "Controlled multi-organization release with explicit gates."])


def add_secure_self_service_and_admin_update(doc):
    doc.add_heading("Authenticated self service and platform administration", 1)
    doc.add_paragraph("Quantum OPS uses ordinary bookmarkable application paths. The public BQ Labs site remains at the root, while every workspace screen verifies an authenticated server session before displaying operational content. Authorization is enforced independently on every API endpoint.")
    add_table(doc, ["Path", "Purpose and access"], [
        ("/quantum-ops/login", "Secure account sign in."),
        ("/quantum-ops/register", "Self-service account creation using mandatory email verification and an optional adaptive human challenge."),
        ("/quantum-ops/setup", "Authenticated setup for workspace, client, team, project, and email-bound invitations."),
        ("/quantum-ops/dashboard", "Role-scoped operating overview."),
        ("/quantum-ops/tasks", "Task queue, creation, assignment, acceptance, and start-work actions."),
        ("/quantum-ops/uat", "Tester assignment, tester acceptance, and observation management."),
        ("/quantum-ops/leave", "Leave and comp-off requests with project impact."),
        ("/quantum-ops/admin", "BQ Labs staff-only administration for account activity, plan usage, and block or restore controls."),
    ], [2.05, 4.45])
    doc.add_heading("Self service operating model", 2)
    add_bullets(doc, [
        "Any genuine user can create an account after email verification, then create one Starter workspace and become its owner.",
        "The owner can create the first client, team, and up to two Starter projects, then invite members, Testers, client viewers, or client approvers by verified email.",
        "Initial Starter safeguards are one client, two projects, five users, 250 active tasks, and 100 MB of attachments. Increasing limits requires BQ Labs approval until payment is implemented.",
        "Project invitations are email-bound, expiring, non-replayable, and grant only the selected project role.",
        "Internal members see only tasks they own, are assigned, or are explicitly testing unless a scoped lead or management role grants a broader view.",
        "Task assignees explicitly accept proposed work before the client sees any state beyond Not Started.",
        "A Task Owner, working member, Team Owner, or Team Lead can nominate Testers. Only an assigned Tester can accept, and only an accepted Tester can raise an observation.",
        "Ordinary members see their own leave records; authorized owners, administrators, management, and relevant Team Leads may see the project-planning view.",
    ])
    doc.add_heading("BQ Labs administration", 2)
    add_bullets(doc, [
        "The Quantum OPS administration page is available only after login and only to active staff accounts controlled by BQ Labs.",
        "It shows registered and active users, last-login activity, task ownership and assignment counts, workspace state, plan, and quota utilization.",
        "A BQ Labs administrator can block or restore another user. The API validates the action and writes an immutable audit event.",
        "Administrators cannot block their own current account, reducing accidental platform lockout risk.",
        "Payment, invoicing, and self-service renewal remain deferred; this page exposes the plan and usage data required for that future phase.",
    ])
    doc.add_heading("Release acceptance additions", 2)
    add_bullets(doc, [
        "Opening a protected friendly URL without a valid session redirects to /quantum-ops/login and returns no workspace data.",
        "A non-staff account receives HTTP 403 from every platform administration endpoint even when the URL is guessed.",
        "A staff administrator can view plan usage and block or restore another account but cannot block themselves.",
        "Member selectors contain only active members of the selected project, with server-side assignment validation.",
        "Tester acceptance and observation creation follow the approved assignment rules and remain protected by API authorization.",
    ])


def build_bqlabs():
    d = Document(); setup(d, "BQ Labs - Product and Brand Brief")
    title(d, "AI-generated strategy brief", "BQ Labs", "Identity, product direction, and public-platform principles", "Approved baseline; update as the company and product portfolio evolve")
    callout(d, "Public naming rule", "Use BQ Labs as the only company name in public-facing copy, interfaces, metadata, and communications.")
    d.add_heading("1. Company proposition", 1)
    d.add_paragraph("BQ Labs creates user-centred software that turns complex operational work into clear, accountable, and useful experiences. Its products should apply AI where it produces measurable user value, while preserving human control, data protection, and explainability.")
    d.add_heading("Vision", 2)
    add_bullets(d, ["Build an AI-enabled organization, not an AI-themed one.", "Create products around genuine user workflows and decisions.", "Design mobile-first, then expand gracefully to larger screens.", "Earn trust through security, reliability, transparency, and restrained public claims."])
    d.add_heading("Brand character", 2)
    add_table(d, ["Attribute", "Expression"], [
        ("Intelligent", "Clear reasoning, useful automation, and evidence-led product decisions."),
        ("Dependable", "Calm interfaces, predictable behaviour, and visible accountability."),
        ("Forward-looking", "Modern interaction patterns without novelty that obstructs work."),
        ("Human", "Plain language, accessible UX, and people remaining in control."),
    ], [1.5, 5.0])
    d.add_heading("2. Product portfolio model", 1)
    d.add_paragraph("The BQ Labs website is the trusted public entry point for a growing product portfolio. Quantum OPS is the first operational product. Future products should inherit shared identity, authentication standards, security posture, design tokens, and support conventions without forcing a single monolithic codebase.")
    add_bullets(d, ["Public site: company story, product discovery, contact paths, and trust signals.", "Product entry: deliberate hand-off from bqlabs.in to the relevant application.", "Shared platform capabilities: identity, organization membership, audit, notifications, and observability.", "Product-specific domains: isolated business logic and permissions for each product."])
    d.add_heading("3. Public experience principles", 1)
    add_bullets(d, ["Lead with outcomes and user problems, not implementation jargon.", "Keep public diagnostics, stack traces, environment details, and infrastructure health private.", "Expose system status only to authenticated staff unless a future public status page is intentionally approved.", "Use fast-loading, responsive pages with semantic HTML, accessible navigation, and minimal client-side payloads.", "Make privacy, support, and security contact paths easy to find before inviting broad adoption."])
    d.add_heading("4. AI principles", 1)
    add_table(d, ["Principle", "Required behaviour"], [
        ("Purpose before AI", "Use AI only when it improves speed, quality, comprehension, or decision support."),
        ("Human authority", "Users review consequential suggestions and retain approval rights."),
        ("Data restraint", "Minimize data shared with models; never treat customer content as public training data."),
        ("Transparency", "Identify AI-generated output, source context, and uncertainty where relevant."),
        ("Safety", "Log model actions, protect prompts and secrets, and test abuse and data-leak scenarios."),
    ], [1.45, 5.05])
    d.add_heading("5. Governance and release discipline", 1)
    add_bullets(d, ["Never store secrets or environment files in Git.", "Never rewrite Git history.", "Request explicit approval immediately before each push to GitHub.", "Request separate explicit approval immediately before each deployment to a server.", "Use staged releases, rollback procedures, database backups, and auditable change records."])
    d.add_paragraph("Planning-document rule: 003.BQLabs-AI-generated.docx and 004.QuantumOPS.docx are synchronized living references. Update both whenever a product, technical, security, scope, or delivery-plan decision changes, even when one document only needs a brief cross-reference or governance update.")
    d.add_paragraph("Client-transparency principle: Quantum OPS supports project-level client visibility policies. BQ Labs may provide a summarized view or a fuller operational view according to the engagement, while mandatory privacy and security exclusions always apply.")
    d.add_paragraph("Team-governance principle: the user who creates a team becomes its Team Owner. Operational Team Lead authority is assigned separately and can be transferred when responsibilities change, preventing access from following a former lead into unrelated projects.")
    d.add_paragraph("Organization-governance principle: the first verified user who creates an organization becomes its Organization Owner and initial Organization Administrator. Organization-level authority can be delegated and recovered without allowing a team role to escalate itself.")
    d.add_paragraph("Initial platform-control principle: any verified team member may request an organization, but BQ Labs must approve the request before tenant resources are provisioned. A dedicated BQ Labs platform-administration console controls approvals, limits, security, and operational governance.")
    d.add_paragraph("Capacity-governance principle: each approved organization starts with bounded, configurable quotas for users, teams, active projects, tasks, and attachments. BQ Labs can raise limits deliberately as usage and infrastructure capacity justify it.")
    d.add_paragraph("Commercial-access principle: Phase 2 remains subscription-ready but does not collect payments. BQ Labs administrators manually control plan, expiry, and access state. Payment gateway, invoicing, refunds, taxation, and self-service renewal are deferred to the following phase.")
    d.add_paragraph("Workforce-operations principle: leave and earned comp-off are included in the Quantum OPS Phase 2 MVP because availability directly affects assignment, forecasts, workload, and delivery risk.")
    d.add_paragraph("Leave-approval principle: leave uses Team Lead approval followed by optional Client approval configured per project. Client participation is limited to project scheduling impact and never reveals private leave reasons or comp-off evidence.")
    d.add_paragraph("Leave-date principle: future and same-day leave follows the configured approval chain, and every affected project that requires client approval must approve. Backdated leave is system-approved automatically after validation and is clearly identified in the audit trail.")
    d.add_paragraph("Leave-revision principle: a client rejection returns a future/same-day leave request as Changes Required rather than terminally rejecting it. Revised dates or handover details create a new auditable revision and restart the approval chain.")
    d.add_paragraph("Leave-response principle: unanswered future/same-day approval steps remain Pending. The system sends reminders and escalation notifications but never converts silence into approval; only validated backdated leave is auto-approved.")
    d.add_paragraph("Leave-reminder default: remind the current approver after 24 hours, remind again after 48 hours, and escalate after 72 hours. Organization Administrators may configure these intervals prospectively, with every policy change audited.")
    d.add_paragraph("Leave-escalation routing: a delayed Team Lead step escalates to the Team Owner. A delayed Client step escalates to the Team Lead, Team Owner, and configured secondary client approver, with role-appropriate data minimization.")
    d.add_paragraph("Client-approval deadline: authorized client approvers may approve only before the leave begins. If any required client step is still pending at the start instant, the request becomes Unapproved Leave and is never auto-approved by the deadline or escalation process.")
    d.add_paragraph("Unapproved-leave balance rule: when an uncancelled request becomes Unapproved Leave, the system deducts or utilizes the balance associated with the requested leave type. Cancellation before the start instant prevents deduction; later corrections require an audited administrator adjustment.")
    d.add_paragraph("Availability-transparency principle: insufficient normal-leave balance never blocks recording an absence. Quantum OPS records the leave and balance shortfall so affected teams and projects can plan; earned comp-off remains unavailable until sufficient balance exists.")
    d.add_paragraph("Shared-calendar principle: organization members may see a colleague's name, leave dates, duration, and status for planning. Leave reason, balance, medical data, comp-off evidence, and administrative notes remain restricted.")
    d.add_paragraph("Task-assignment principle: a Team Lead proposes an assignment, but the team member must explicitly act. Assignment is not treated as accepted work until the member accepts or raises a structured reassignment, ETA, or clarification request.")
    d.add_paragraph("Project-calendar principle: assignment response deadlines use the next working day from the task's project calendar. Workweek, timezone, holidays, exceptional working days, and the project leave register are configurable per project, with controlled Excel import for bulk leave data.")
    d.add_paragraph("Assignment-overdue principle: when the next-working-day deadline passes without member action, the assignment remains unresolved as Pending Member Action - Overdue and escalates to the Team Lead. Silence never creates acceptance.")
    d.add_paragraph("Assignment-recovery principle: the Team Lead may withdraw or reassign an overdue proposal without the original member's response. A reason is mandatory, the prior proposal closes explicitly, and all assignment revisions remain auditable.")
    d.add_paragraph("Client-status principle: until a proposed team member accepts, the client-facing task status is always Not Started. Internal proposal, overdue, assignee, withdrawal, and reassignment details remain hidden even under Full Detail visibility.")
    d.add_paragraph("Client-assignee principle: after acceptance, Summarized projects show only the responsible team, while Full Detail projects may show the accepted assignee's name. Pre-acceptance identity remains hidden under both profiles.")
    d.add_paragraph("Task-start principle: member acceptance confirms ownership but does not imply execution. The task remains Not Started until the accepted assignee explicitly selects Start Work, which moves it to In Progress.")
    d.add_paragraph("Missed-start principle: passing the planned start without Start Work does not mutate task status. The task remains Not Started while a separate start-delay alert notifies the member and Team Lead and is recorded for reporting.")
    d.add_paragraph("Missed-start reminder principle: while the start-delay alert remains open, notify the member and Team Lead once per project working day. Stop when work starts, planned start changes, assignment changes, or the task is cancelled.")
    d.add_paragraph("Progress-update principle: once work is In Progress, daily updates are not mandatory. The member may update status or add progress when applicable; the system records real changes without manufacturing compliance activity.")
    d.add_paragraph("Task-status principle: the assigned team member may update every task status through valid transitions. Completed reviews create separate review records including Reviewed By, timestamp, outcome, and optional notes.")
    d.add_paragraph("Design-system principle: BQ Labs, Quantum OPS, and future products use one reusable Apple-inspired visual foundation: clarity, generous space, premium typography, restrained motion, rich product storytelling, and accessibility. It must remain distinctly BQ Labs and never copy Apple's assets, marks, or proprietary page compositions.")
    d.add_paragraph("UAT-authorship principle: a UAT observation may be raised by the task's assigned team member or by a project member explicitly assigned as Tester for that task. A general Team Lead or client role does not grant observation-authoring rights by itself.")
    d.add_paragraph("Task-review principle: reviewer independence is not mandatory. Any authorized project participant, including the assignee, another member, Team Lead, or client with review permission, may review; the system records identity and role for transparency.")
    d.add_paragraph("Review-gate principle: review is configurable per task. Tasks marked Review Required cannot enter UAT without a completed review for the current task revision; tasks marked Review Optional may proceed without one.")
    d.add_paragraph("Task-ownership principle: the user who creates a task becomes its Task Owner. Only that Task Owner may change the task's Review Required setting, unless task ownership itself is formally transferred through an audited workflow.")
    d.add_paragraph("Team-Lead succession principle: an Organization Administrator marking a Team Lead as resigned must select a replacement. The system atomically transfers lead authority, active task ownership, pending approvals, escalations, and management duties while preserving all historical attribution.")
    d.add_heading("6. BQ Labs experience system", 1)
    d.add_paragraph("Create a shared design system that gives the public BQ Labs site, Quantum OPS, and future products a coherent family resemblance while allowing each product an appropriate operational personality. The inspiration is Apple's emphasis on short benefit-led messages, large visual moments, simple navigation, product-family consistency, strong privacy communication, and fluid cross-device experiences.")
    add_table(d, ["Design quality", "BQ Labs interpretation"], [
        ("Clarity", "One primary message per section; plain language; strong visual and task hierarchy."),
        ("Space", "Generous breathing room on public pages; disciplined density and progressive disclosure inside applications."),
        ("Typography", "Large editorial display type for storytelling, highly readable interface type for operational work."),
        ("Imagery", "Original BQ Labs product renders, interface compositions, diagrams, and photography; no copied Apple assets."),
        ("Motion", "Purposeful scroll reveals, state transitions, and feedback using transform/opacity; reduced-motion support is mandatory."),
        ("Consistency", "Shared tokens, components, icon rules, navigation, accessibility, and interaction conventions across products."),
        ("Trust", "Privacy, security, system boundaries, and data controls explained as clearly as product benefits."),
    ], [1.45, 5.05])
    add_bullets(d, ["Public BQ Labs pages may be cinematic and editorial, with large product showcases and restrained calls to action.", "Quantum OPS should feel premium but operational: calm surfaces, dense information only where useful, strong tables/lists, responsive dashboards, and minimal decoration around serious work.", "Future products inherit core design tokens and shell components, then add a small product accent palette and product-specific components.", "Performance is part of the design: responsive images, minimal JavaScript, no blocking animation, stable layout, and meaningful content without motion.", "The final design phase will review moodboards, typography, color, page architecture, component density, dashboard patterns, imagery direction, and motion prototypes before implementation."])
    d.add_heading("7. Success measures", 1)
    add_bullets(d, ["Visitors understand the company and first product within one minute.", "Core public pages meet agreed accessibility and performance budgets.", "Product discovery converts to legitimate sign-in, demo, or contact intent.", "No operational diagnostics or sensitive metadata are publicly exposed.", "Brand and interaction patterns remain coherent as new products are introduced."])
    d.add_heading("8. Open decisions", 1)
    add_bullets(d, ["Final public positioning statement and primary audience.", "Product naming and subdomain convention.", "Support, privacy, and legal publication readiness.", "AI governance owner and approved model/provider policy.", "Company-registration and claims review before public commercialization."])
    add_final_design_blueprint(d, compact=True)
    add_secure_self_service_and_admin_update(d)
    d.save(ROOT / "003.BQLabs-AI-generated.docx")


def build_quantumops():
    d = Document(); setup(d, "Quantum OPS - Phase 2 Product and Technical Plan")
    title(d, "Product requirements and implementation plan", "Quantum OPS", "Multi-team delivery operations with a secure client portal", "Planning baseline - client login included in MVP")
    callout(d, "Product decision", "Quantum OPS will support authenticated clients in the MVP. A client can see tasks only within projects explicitly assigned to that client. Each project selects a configurable client-visibility profile: Summarized or Full Detail. New organizations require BQ Labs platform approval before provisioning.")
    d.add_heading("1. Product intent and operating model", 1)
    d.add_paragraph("Quantum OPS is a multi-tenant delivery-operations platform for BQ Labs teams and their clients. It combines project visibility, task execution, delivery risk, UAT resolution, capacity awareness, and management reporting while keeping each organization's and project's information isolated.")
    d.add_heading("Primary outcomes", 2)
    add_bullets(d, ["Team members always know what to do next and what is blocked.", "Team leads can plan, assign, approve, and intervene before delivery slips.", "Clients receive trustworthy project visibility and can act on UAT or information requests.", "Senior management sees portfolio health without inspecting every task.", "Every consequential change can be attributed and reconstructed."])
    d.add_heading("2. Users and roles", 1)
    add_table(d, ["Role", "Primary capabilities", "Default boundary"], [
        ("BQ Labs platform admin", "Approve organizations; manage tenant state, quotas, security events, support access, and platform operations.", "Whole platform metadata; tenant content only through audited support access."),
        ("Team member", "Work assigned tasks; add updates, blockers, comments, and requests.", "Assigned and team-visible work."),
        ("Team owner", "Retain team-level governance, assign or replace leads, and manage project visibility.", "The created team; ownership is continuous until transferred."),
        ("Team lead", "Plan projects, assign work, approve requests, manage capacity, risk, and project visibility.", "Currently assigned teams and projects only."),
        ("Client", "View project tasks and approved project information; comment; submit/verify UAT; respond to requests.", "Tasks belonging to explicitly granted client projects only."),
        ("Senior management", "View cross-team portfolio, trends, capacity, and escalations.", "Organization-wide read access; limited mutation."),
        ("Organization owner", "Create the organization, appoint administrators, transfer ownership, and provide governance continuity.", "One organization; exactly one active owner."),
        ("Organization admin", "Manage policies, teams, membership, administrator assignments, ownership recovery, and security settings.", "One organization; appointed by owner or authorized admin."),
    ], [1.1, 3.4, 2.0])
    d.add_paragraph("Roles are assignments within a scope, not permanent global flags. A person may lead one team, contribute to another, and hold client access in a separate organization.")
    d.add_heading("3. Domain model", 1)
    add_table(d, ["Area", "Core entities"], [
        ("Tenancy", "OrganizationRequest, Organization, OrganizationOwnership, User, Membership, RoleAssignment, AdminAssignment, Team, TeamOwnership, TeamMembership, LeadAssignment, LeadSuccession"),
        ("Client access", "ClientAccount, ClientContact, ProjectMembership, ProjectVisibilityPolicy, ProjectLeaveApprovalPolicy, VisibilityPolicyRevision"),
        ("Delivery", "Project, ProjectCalendar, ProjectWorkingDay, ProjectCalendarException, Module, Milestone, Release, Task, TaskDependency, TaskAssignment, AssignmentDecision, TaskTesterAssignment, TaskReview"),
        ("Execution", "WorkUpdate, Comment, Mention, Attachment, TaskStatusEvent"),
        ("Control", "ChangeRequest, ApprovalDecision, Blocker, UATObservation"),
        ("People operations", "LeavePolicy, WorkingCalendar, Holiday, ProjectLeaveImport, ProjectLeaveImportRow, LeaveRequest, LeaveProjectImpact, LeaveApprovalStep, LeaveDecision, CompOffClaim, CompOffWorkEntry, CompOffLedgerEntry, AvailabilityEvent"),
        ("Platform", "Notification, NotificationPreference, AuditEvent, OutboxEvent"),
    ], [1.45, 5.05])
    d.add_heading("Task definition", 2)
    add_bullets(d, ["Stable human identifier, such as QOPS-1042, plus an internal UUID.", "The creating user is stored as Task Owner separately from assignee, tester, reporter, reviewer, and Team Lead.", "Project, module, type, description, acceptance criteria, priority, complexity (1-5), expected effort, confidence, and per-task Review Required flag.", "Planned start, planned finish, forecast finish, and actual completion instead of one ambiguous ETA.", "Assignee, optional task tester assignment, collaborators, reporter, optional reviewer, dependencies, client visibility, and release or milestone.", "Only Task Owner can change Review Required; every change records prior value, new value, actor, timestamp, and task revision.", "Review records contain Reviewed By, reviewer role, reviewed timestamp, outcome, optional notes, task revision, and evidence references.", "Reviewer may be the assignee, another authorized team member, Team Lead, or authorized client; review permission remains project-scoped.", "A completed review satisfies the gate only for the reviewed task revision. Material changes to description, acceptance criteria, implementation evidence, or scope mark the review stale and require a new review when Review Required is enabled.", "Internal status and client-facing status are mapped explicitly; client APIs never infer or expose internal workflow states accidentally.", "Current status is optimized for reads; every transition is also stored as an immutable event."])
    d.add_heading("4. Core workflows", 1)
    d.add_heading("Project delivery", 2)
    add_steps(d, ["Create a project, its timezone and working calendar, modules, milestones, members, and client visibility policy.", "Create and estimate work; validate acceptance criteria, dependencies, and whether this task requires review.", "The Team Lead proposes an assignee; the assignment enters Pending Member Action and notifies the member.", "The response deadline is the end of the next project working day, skipping configured non-working weekdays, holidays, and non-working exceptions.", "The member accepts, requests reassignment, proposes an ETA revision, or requests description/acceptance-criteria clarification.", "If the deadline passes without action, the assignment becomes Pending Member Action - Overdue, remains unaccepted, notifies the Team Lead, and appears in lead and member dashboards.", "The Team Lead may withdraw the overdue proposal or propose a different member. The action requires a reason, closes the prior proposal, notifies the affected members, and creates a new assignment revision with its own response deadline.", "Acceptance activates ownership while task execution remains Not Started. A request routes back to the Team Lead for a recorded decision and leaves the assignment unresolved until the workflow is completed.", "If the planned start passes first, keep status Not Started, create a separate start-delay alert, and notify the accepted member and Team Lead without inventing another task status.", "While unresolved, send one reminder on each subsequent project working day. Deduplicate retries so a member or lead receives no more than one scheduled missed-start reminder per task per working day.", "The accepted assignee explicitly selects Start Work; the system timestamps the action, closes the start-delay alert if present, and transitions the task from Not Started to In Progress. Changing planned start, assignment, or task cancellation also closes/recalculates the alert as appropriate.", "The assigned member may perform every permitted status transition and may add progress, comments, evidence, blockers, and forecast changes whenever applicable; no daily update is required.", "When review is performed, create a TaskReview record with reviewer identity, timestamp, outcome, notes, evidence, and reviewed task revision.", "Before UAT, enforce the task's Review Required flag. A current completed review satisfies the gate; an absent or stale review blocks the transition with a clear explanation. Review Optional tasks may proceed.", "Close tasks and milestones; preserve the timeline for reporting and audit."])
    d.add_heading("Task lifecycle", 2)
    d.add_paragraph("Assignment workflow: Proposed -> Pending Member Action -> Accepted or Closed/Reassigned. Execution workflow: Not Started -> In Progress -> Blocked -> Review -> UAT -> Completed, with Cancelled available through an authorized terminal transition. Acceptance does not change execution status; Start Work does. Hold is captured as a pause reason rather than an ordinary delivery stage. Client mapping: every pre-acceptance and accepted-but-not-started state appears as Not Started.")
    d.add_heading("Structured change requests", 2)
    add_bullets(d, ["Reassignment request", "Forecast/ETA change request", "Description or acceptance-criteria enhancement request", "Scope or priority change request"])
    d.add_paragraph("Each request records the assignment revision, proposal, reason, requester, Team Lead decision, comment, timestamps, and resulting task revision. Silence never counts as member acceptance.")
    d.add_heading("Blockers and escalation", 2)
    add_bullets(d, ["Description, category, severity, owner, helper or resolver, expected resolution date, and dependency.", "Age-based escalation and explicit resolution notes.", "Client-visible only when an authorized internal user deliberately publishes it."])
    d.add_heading("UAT observations", 2)
    add_bullets(d, ["Only the assigned team member or the project member explicitly assigned as Tester for the task may raise an observation.", "Tester authority is task-scoped, starts and ends with the TaskTesterAssignment, and is enforced server-side.", "Reporter, observation, reproduction steps, expected result, actual result, environment, evidence, severity, owner, and status.", "Fix description and type: understanding issue, new requirement, code fix, or data fix.", "Clients and Team Leads may view or participate according to project visibility, but cannot create a UAT observation unless they also hold an allowed task assignment.", "Verification and closure are immutable events linked to the applicable task and observation revision."])
    d.add_heading("Leave and comp-off", 2)
    d.add_paragraph("Leave and earned comp-off are required Phase 2 MVP workflows. They remain a bounded module but integrate with team availability, assignment warnings, capacity dashboards, notifications, and audit reporting.")
    add_steps(d, ["A team member submits leave dates, leave type, duration, private reason, optional handover notes, and the projects affected by the absence.", "The system checks date validity, overlaps, policy constraints, organization time zone, and current task commitments, then derives a project-impact summary.", "If the leave start date is earlier than the organization's current local date, the system classifies it as Backdated and auto-approves it after validation, recording a System decision and audit event.", "If the leave starts on the current local date or later, the currently assigned Team Lead performs mandatory Level 1 approval and records approve, reject, or request-clarification with a reason.", "After Level 1 approval, every affected project whose ProjectLeaveApprovalPolicy requires client approval creates a separate Level 2 step for an authorized client approver.", "All required client project steps must approve before the leave starts. One client's approval applies only to that client's project and cannot approve another project's impact.", "If any required client declines before the deadline, the request becomes Changes Required and records that project's safe client-facing reason; it is not terminally rejected.", "The member and Team Lead may create a revision with changed dates, project impact, coverage, or handover. Resubmission restarts Level 1 and all applicable Level 2 steps.", "When no affected project requires client approval, Team Lead approval is final. Otherwise the leave becomes fully approved only after all required Level 2 steps approve the current revision before the start instant.", "At the leave start instant, any still-pending required client step closes without approval and the overall request becomes Unapproved Leave. No client or internal approver may retroactively approve that future-dated revision afterward.", "If the request was not cancelled before the start instant, the system atomically posts the appropriate leave deduction or comp-off Utilized ledger entry and creates availability and delivery-risk events.", "Cancellation before the leave begins closes the request as Cancelled and releases reservations without deduction. Cancellation after the start instant cannot reverse balance automatically; an authorized administrator must post a reasoned adjustment.", "Final or automatic approval creates availability events and warns leads about affected assignments and milestones without automatically reassigning work.", "Cancellation or modification always preserves the prior request and applicable approval history rather than silently rewriting it."])
    d.add_heading("Comp-off earning and use", 3)
    add_steps(d, ["A member submits earned-comp-off evidence as one or more work entries: date, related task, hours worked, and reason.", "The approver validates the evidence and approves the number of earned hours or days.", "Approval posts an immutable Earned entry to the member's comp-off ledger.", "A comp-off leave request reserves available balance while pending and posts a Utilized entry only when approved.", "Rejection or cancellation releases the reservation; expiry and administrative adjustment use separate ledger entries with reasons."])
    add_bullets(d, ["A member cannot reserve or utilize more comp-off than the available approved balance.", "Normal leave is still recorded when balance is insufficient; the ledger reports the resulting shortfall without blocking the availability event.", "Balances are derived from immutable ledger entries, not stored as a freely editable number.", "Dates, partial days, working calendars, holidays, time zones, overlaps, and decimal-hour precision require explicit validation.", "All organization members may view name, leave dates, duration, and current status on the shared availability calendar.", "Leave reason, normal-leave balance, comp-off balance/evidence, medical information, administrative adjustments, and internal approval notes are excluded from the shared calendar.", "Client approval is configurable independently for each project and names primary and optional secondary client approvers authorized to decide.", "Client users see only leave dates, duration, affected project, delivery impact, and approved handover summary. They never see private reasons, leave balance, medical information, comp-off evidence, or unrelated projects.", "Each approval step has its own state, actor, timestamp, comment, revision number, and immutable audit event; a client cannot approve before Level 1 approval or on/after the leave start instant.", "Approvals from an older revision never carry forward automatically, because changed dates or coverage may alter every project's decision.", "Pending future/same-day approval steps remind at 24 hours, remind again at 48 hours, and escalate at 72 hours by default, but never auto-approve because of elapsed time.", "A pending Team Lead step escalates to the Team Owner. A pending Client step escalates to the Team Lead, Team Owner, and configured secondary client approver.", "Escalation messages are authorization-aware: internal recipients may see permitted operational context, while client recipients receive only safe project-impact information.", "Organization Administrators may configure reminder intervals; new settings apply prospectively and every policy revision is audited.", "The server transitions overdue pending approval steps to Unapproved Leave at the leave start instant using the organization's configured time zone and records an audit event.", "Unapproved Leave deducts the requested normal leave balance, including recording a shortfall, or utilizes reserved comp-off balance unless cancelled before the start instant.", "Unapproved Leave is visible to the member, Team Lead, Team Owner, and authorized Organization Administrators; clients see only the resulting project availability impact.", "Backdated auto-approval bypasses Team Lead and Client decisions but never bypasses validation, authorization to submit, comp-off balance rules, overlap checks, or audit logging.", "The server, not the browser, calculates whether leave is backdated using the organization's configured time zone."])
    d.add_heading("Project calendars and Excel leave import", 3)
    add_bullets(d, ["Each project defines timezone, normal working weekdays, working-day start/end, holidays, non-working exceptions, and optional exceptional working days.", "Calendar revisions are effective-dated and audited; existing deadlines retain the calendar version used when they were calculated unless explicitly recalculated.", "Team Owner and current Team Lead may maintain the project calendar. Organization Administrators may provide organization defaults, but a project may override them.", "The project leave register is a view over canonical leave records linked to that project, including imported records; it is not a second source of truth.", "Excel import uses a preview-and-confirm workflow. No records are committed until every row is validated or the authorized importer explicitly excludes invalid rows."])
    add_table(d, ["Excel column", "Required", "Validation/use"], [
        ("Employee Email", "Yes", "Must match an active organization member assigned to the project."),
        ("From Date", "Yes", "ISO date preferred; interpreted in the project timezone."),
        ("To Date", "Yes", "Must be on or after From Date."),
        ("Duration", "Yes", "Full Day, First Half, Second Half, or valid hours."),
        ("Leave Type", "Yes", "Must map to an enabled organization leave type or Comp Off."),
        ("Status", "Yes", "Approved, Unapproved Leave, Cancelled, or Pending Import Review."),
        ("Reason", "No", "Private; excluded from project calendar and all client views."),
        ("External Reference", "No", "Used for idempotent re-import and duplicate prevention."),
    ], [1.45, .75, 4.3])
    add_bullets(d, ["Imports validate file type, size, headers, formulas, dates, membership, project scope, duplicates, overlapping leave, and comp-off balance.", "Imported future leave cannot bypass the agreed Team Lead and optional Client approval workflow unless its status is an explicitly authorized historical record.", "Backdated imported leave follows the backdated auto-approval rules after validation.", "Every import stores uploader, source filename, checksum, row counts, errors, timestamp, and created record IDs; original files are private and retention-controlled."])
    d.add_heading("5. Permissions and client portal", 1)
    d.add_heading("BQ Labs platform administration", 2)
    add_steps(d, ["Any email-verified team member submits an organization request with name, purpose, expected team size, and expected usage.", "The request enters Pending Review; no organization database scope, teams, projects, or client invitations are provisioned yet.", "A BQ Labs Platform Administrator approves, rejects, or requests more information and records the reason.", "On approval, the system atomically provisions the organization and makes the requester its Organization Owner and initial Organization Administrator.", "The requester is notified; all request and provisioning events are retained in the platform audit trail."])
    add_bullets(d, ["Platform Administrators are distinct from Organization Administrators and are provisioned only by controlled BQ Labs operations, never through public self-promotion.", "The platform console shows pending approvals, active/suspended organizations, user and project counts, storage and notification usage, quota warnings, failed jobs, security alerts, and recent admin actions.", "Admin actions include approve/reject, adjust quotas, suspend/reactivate an organization, revoke sessions, inspect delivery failures, and initiate explicitly authorized support access.", "Routine platform administration exposes operational metadata rather than tenant task content. Time-bound support access to tenant content requires a reason, re-authentication, prominent audit logging, and notification according to policy.", "All admin lists are paginated, searchable, filterable, and backed by aggregated usage data so the dashboard does not overload the application."])
    d.add_heading("Starter organization quotas", 3)
    add_table(d, ["Resource", "Initial limit", "Enforcement"], [
        ("Active users", "25", "Block additional activation or invitation acceptance until capacity is released or raised."),
        ("Teams", "5", "Block new team creation at the limit."),
        ("Active projects", "10", "Allow archival; block creation or reactivation beyond the limit."),
        ("Tasks", "10,000", "Count non-deleted organization tasks; block creation at the hard limit."),
        ("Attachments", "2 GB", "Block new uploads at the hard limit while preserving access to existing files."),
    ], [1.35, 1.05, 4.1])
    add_bullets(d, ["Show dashboard warnings at 80% and 95% consumption and notify organization administrators.", "Never delete or corrupt existing data when a quota is reached; restrict only the action that increases consumption.", "BQ Labs Platform Administrators may raise or lower individual limits with a reason and immutable audit event.", "A reduced limit below current usage prevents further growth but does not remove existing data.", "Quota checks execute atomically on the server to prevent concurrent requests from exceeding the approved limit."])
    d.add_heading("Organization and team authority", 2)
    add_bullets(d, ["After BQ Labs approves the request, the verified requester becomes its Organization Owner and initial Organization Administrator in one atomic provisioning transaction.", "The Organization Owner and authorized Organization Administrators may appoint or remove additional Organization Administrators.", "Organization Owner and Organization Administrators may transfer Team Ownership; the current Team Owner may also transfer ownership of their own team.", "A Team Owner or Team Lead cannot promote themselves to an organization or platform role merely through team permissions.", "The last Organization Owner cannot be removed, deactivated, or leave until ownership is transferred to another verified organization member.", "All administrator appointments, removals, ownership transfers, and emergency recovery actions require re-authentication and immutable audit events."])
    d.add_heading("Team Lead resignation and succession", 3)
    add_steps(d, ["An Organization Administrator opens Mark Team Lead Resigned and selects the affected team/project scope.", "The administrator must select an active, eligible replacement before confirmation; the operation cannot leave the scope without a Team Lead.", "The system previews active tasks owned by the departing lead, tasks personally assigned to the departing lead, pending approvals, assignment escalations, client-visibility controls, leave decisions, alerts, and other management responsibilities that will transfer.", "After re-authentication and confirmation, one transaction ends the former LeadAssignment, creates the replacement LeadAssignment, transfers active Task Ownership and open management work, and records a LeadSuccession event.", "Each active task personally assigned to the departing lead creates a new proposal to the replacement as Pending Member Action; it is not silently accepted. The replacement follows the normal next-working-day response workflow.", "The replacement is notified and receives a succession dashboard summarizing transferred responsibilities; affected internal stakeholders receive a role-change notification.", "Completed, cancelled, and archived tasks retain the original Task Owner. Historical comments, decisions, reviews, and audit events always retain the actor who performed them."])
    add_bullets(d, ["Active means any non-terminal task or open approval/work item in the departing lead's authorized scope.", "Assignments already accepted by other team members do not change; only ownership and management responsibility transfer.", "Personally executed tasks use Pending Member Action for the replacement so the previously approved assignment-acceptance rule remains intact.", "The departing lead loses lead-only authorization immediately when the transaction commits, and existing sessions must refresh their permissions.", "The operation is idempotent, auditable, and rolls back completely if any required transfer fails."])
    add_table(d, ["Action", "Member", "Lead", "Client", "Management"], [
        ("View internal task details", "Scoped", "Yes", "No", "Read"),
        ("View task in an assigned client project", "Scoped", "Yes", "Yes", "Read"),
        ("Change delivery status", "All valid transitions on assigned task", "Yes", "No", "No"),
        ("Accept or question proposed assignment", "Own proposal", "Resolve request", "No", "No"),
        ("Submit UAT observation", "Assignee or assigned Tester", "Only if task-assigned", "Only if task-assigned", "No"),
        ("Verify UAT resolution", "No", "Yes", "Yes", "No"),
        ("Record task review", "If project-authorized", "Yes", "If project-authorized", "No"),
        ("Change Review Required", "Task Owner only", "Only if Task Owner", "No", "No"),
        ("Approve leave - Level 1", "No", "Yes", "No", "Override only"),
        ("Approve leave - Level 2", "No", "No", "If project requires", "No"),
        ("Manage project access", "No", "Yes", "No", "Limited"),
        ("View audit trail", "Own/team subset", "Project", "Own actions", "Organization"),
    ], [2.15, 1.0, .9, 1.0, 1.45])
    d.add_heading("Project-level visibility profiles", 2)
    add_table(d, ["Profile", "Client experience", "Typical use"], [
        ("Summarized", "Milestones, selected tasks, mapped client status, responsible team, approved ETA, published blockers, decisions, and UAT.", "Clients preferring concise governance reporting."),
        ("Full Detail", "All project tasks and permitted operational fields after acceptance, including accepted assignee name, estimates, work updates, dependencies, and client-visible discussion.", "Highly collaborative or embedded client teams."),
    ], [1.15, 3.65, 1.7])
    add_bullets(d, ["Summarized is the safe default for every new project.", "The Team Owner (the team creator) and the currently assigned Team Lead may change the profile only with an explicit reason; the change and affected fields are audited.", "Team Lead is a transferable, scoped assignment. Removing or replacing the lead immediately revokes their lead-only authority for the former team or project.", "A newly assigned lead receives visibility-management authority only for the scope explicitly assigned to them.", "Optional field controls can refine a profile per project, but no individual task can expand beyond the project's maximum policy.", "Regardless of profile, a client sees Not Started until member acceptance and cannot see the proposed assignee, Pending Member Action, overdue, withdrawal, or reassignment history.", "After acceptance, Summarized returns the responsible team but not the individual; Full Detail may return the accepted assignee's display name.", "Full Detail never exposes secrets, security diagnostics, credentials, private HR or leave information, privileged audit/security data, or information from another project.", "A client may query or open tasks only when the task belongs to a project explicitly assigned to that client account or contact.", "Project scope and field visibility are enforced by backend authorization and tenant-aware serializers/queries, not merely hidden navigation.", "Object identifiers alone never grant access, and cross-project access attempts fail closed and are logged.", "Invitations are expiring, revocable, hashed at rest, and bound to intended email and scope."])
    d.add_heading("6. Dashboards", 1)
    add_table(d, ["Audience", "Default dashboard"], [
        ("BQ Labs platform admin", "Organization approvals, tenant health, quotas, usage trends, failed jobs, security alerts, support sessions, and admin audit."),
        ("Team member", "My work, pending/overdue assignment actions, missed-start alerts, due soon, blockers, mentions, leave/comp-off status and balance, pending requests, and team availability."),
        ("Team lead", "Delivery health, overdue member actions, missed-start alerts, unassigned work, workload, cross-team availability calendar, upcoming leave, approvals, forecast risk, blockers, UAT trends."),
        ("Client", "Milestones, approved progress, project leave-impact decisions when enabled, UAT items, published risks, recent updates."),
        ("Senior management", "Portfolio health, schedule confidence, capacity, aging risks, delivery trends, intervention list."),
    ], [1.5, 5.0])
    d.add_paragraph("Health is calculated from objective signals—forecast variance, overdue dependencies, blocker severity and age, UAT trends, and update freshness—rather than manually selected colours.")
    d.add_heading("7. Notifications", 1)
    add_bullets(d, ["In-app notification centre in MVP; email for invitations, high-priority actions, and user-selected digests.", "Events: assignment, missed planned start, mention, due/forecast change, blocker escalation, leave or comp-off submission/decision/cancellation/reminder/escalation, approval, UAT activity, access change, and milestone update.", "Per-user preferences, deduplication, read state, digest windows, and quiet periods.", "Outbox pattern and asynchronous workers prevent notification delivery from slowing requests."])
    d.add_heading("8. Auditability, security, and privacy", 1)
    add_bullets(d, ["Immutable audit events for membership, roles, client visibility, task transitions, approvals, exports, and sensitive edits.", "Strong authentication; MFA-ready design; session revocation; secure password reset; rate limiting and suspicious-login telemetry.", "Server-side input validation, output encoding, CSRF protection, safe file handling, and narrowly configured CORS.", "PostgreSQL constraints reinforce tenant integrity and legal state transitions.", "Soft deletion or archival for business records, with retention and privacy policies agreed before launch.", "No public diagnostics, stack traces, secrets, infrastructure topology, or internal system status."])
    d.add_heading("9. Responsive UX and accessibility", 1)
    add_bullets(d, ["Mobile (360px+): My Work, quick updates, blocker creation, comments, UAT response, and approvals.", "Tablet: split list/detail and compact planning views.", "Desktop: portfolio, planning board, capacity, advanced filters, bulk actions, and reporting.", "List/detail is the primary interaction; Kanban is an optional view rather than the only way to work.", "Keyboard navigation, visible focus, semantic headings, labelled controls, accessible errors, adequate contrast, and 44px touch targets."])
    d.add_heading("Apple-inspired, BQ Labs-owned design direction", 2)
    d.add_paragraph("Quantum OPS will share the BQ Labs experience system while translating its premium, spacious, product-focused character into a serious operational application. The objective is Apple-like confidence and craft, not imitation.")
    add_table(d, ["Surface", "Design direction"], [
        ("BQ Labs public site", "Editorial product storytelling, bold benefit headlines, original product imagery, generous space, restrained calls to action, and smooth section transitions."),
        ("Quantum OPS sign-in/onboarding", "Focused single-purpose screens, calm brand moments, clear progress, and minimal cognitive load."),
        ("Quantum OPS application", "Quiet navigation shell, crisp hierarchy, progressive disclosure, high-quality tables/lists, clear states, and fast keyboard/touch workflows."),
        ("Dashboards", "Strong narrative summary at the top, then actionable detail; meaningful color only for status, risk, and decisions."),
        ("Future BQ Labs products", "Shared tokens and components with a controlled product accent and domain-specific patterns."),
    ], [1.65, 4.85])
    add_bullets(d, ["Create original BQ Labs typography, color, iconography, product imagery, and motion language; never use Apple's logos, product images, copy, or cloned page sections.", "Use a reusable token system for color, type scale, spacing, radii, shadows, motion duration/easing, breakpoints, density, focus, and semantic states.", "Support light and dark themes where they improve the product, with accessible contrast and user/system preference.", "Motion remains optional and purposeful; reduced-motion mode delivers the same information and task completion path.", "Public storytelling may use richer motion and imagery, while authenticated operational screens prioritize speed, stability, clarity, and lower bundle cost.", "The dedicated end-of-planning design discussion will produce a moodboard, page map, responsive wireframes, design tokens, component inventory, motion rules, and prototype acceptance criteria before implementation."])
    d.add_heading("10. Technical architecture", 1)
    add_table(d, ["Layer", "Recommendation"], [
        ("Frontend", "React with route-based code splitting, query caching, accessible component primitives, and mobile-first CSS."),
        ("API", "Django + Django REST Framework; versioned endpoints; serializers as validation boundary; policy services for authorization."),
        ("Database", "PostgreSQL with UUID keys, tenant-aware indexes, constraints, transactions, and migration discipline."),
        ("Async", "Redis-backed worker for email, digests, exports, and escalations; transactional outbox for reliability."),
        ("Files", "Private object storage with short-lived signed access and malware/type/size validation."),
        ("Operations", "Structured logs, metrics, traces, private health checks, backups, restore tests, and rollback playbooks."),
    ], [1.25, 5.25])
    d.add_heading("Performance and scale budgets", 2)
    add_bullets(d, ["Paginate and filter all collections server-side; enforce bounded page sizes.", "Use select_related/prefetch_related and query-count tests for dashboard endpoints.", "Cache expensive aggregates with tenant-aware keys and explicit invalidation.", "Define MVP targets: p95 read API under 500 ms, p95 write under 800 ms, and responsive interaction feedback under 100 ms where local.", "Lazy-load secondary views and keep initial JavaScript and API payloads within agreed budgets."])
    d.add_heading("11. Subscription readiness and future payment phase", 1)
    d.add_paragraph("Phase 2 defines organization-level plans, entitlements, expiry, and restricted-access behaviour but does not process money. Prices remain planning assumptions in Indian rupees, exclusive of applicable taxes, and require commercial validation before the future payment phase.")
    add_table(d, ["Plan", "Planning price", "Included capacity and use"], [
        ("Pilot", "₹0 for 30 days", "10 internal users, 2 teams, 2 active projects, 1,000 tasks, 500 MB; client portal included."),
        ("Team", "₹4,999/month or ₹49,990/year", "25 internal users, 5 teams, 10 active projects, 10,000 tasks, 2 GB; client accounts included."),
        ("Business", "₹9,999/month or ₹99,990/year", "50 internal users, 15 teams, 30 active projects, 50,000 tasks, 10 GB; priority support and advanced portfolio controls."),
        ("Enterprise", "Custom annual quote", "Negotiated capacity, SSO, enhanced support/SLA, security review, data controls, and optional dedicated deployment."),
    ], [1.05, 1.65, 3.8])
    add_bullets(d, ["Client accounts associated with authorized projects do not consume internal-user seats at launch; fair-use and abuse controls still apply.", "Plan and price definitions are versioned so future billing can preserve contracted terms.", "Phase 2 stores plan, entitlement snapshot, start date, expiry date, status, and administrator reason without collecting card, bank, or payment-gateway data.", "BQ Labs Platform Administrators can activate, extend, expire, or suspend an organization manually, with re-authentication and an immutable audit event."])
    d.add_heading("Read-only subscription state", 2)
    add_bullets(d, ["When the administrator-set subscription expires or is placed in payment-related suspension, authenticated users retain authorized read-only access to existing data.", "All operational writes are blocked, including task creation or assignment, status changes, comments, uploads, invitations, approvals, settings, and exports that create background artifacts.", "Phase 2 shows plan, expiry, restricted-state reason, and a contact-BQ-Labs renewal path; it does not accept payment.", "A BQ Labs Platform Administrator restores or extends access after the external commercial process is complete.", "A separate Security Locked state removes ordinary access completely and cannot be bypassed through a plan extension.", "Show expiry reminders and a persistent restricted-mode banner; never silently discard data."])
    d.add_heading("Deferred payment scope", 2)
    add_bullets(d, ["Payment gateway and webhook processing", "Invoices, receipts, taxes, credits, and refunds", "Online checkout and stored billing methods", "Self-service renewal, upgrade, downgrade, and cancellation", "Failed-payment retries and automated dunning", "Financial reconciliation and payment reporting"])
    d.add_heading("12. MVP scope", 1)
    add_table(d, ["MVP", "Following increment"], [
        ("BQ Labs platform admin console, organization requests, approvals, tenant state, quotas, and admin audit", "Automated plan tiers, billing, and advanced abuse detection"),
        ("Subscription-ready plan/entitlement state, admin-controlled expiry, and server-enforced read-only restrictions", "Payment gateway, invoicing, taxes, refunds, reconciliation, and self-service renewal"),
        ("Authentication; organizations; teams; invitations; scoped roles", "Configurable enterprise identity and SSO"),
        ("Projects, modules, milestones, tasks, dependencies, assignments", "Templates, recurring operations, and advanced planning"),
        ("Comments, updates, blockers, change requests, audit trail", "Advanced workflow customization and analytics"),
        ("Leave requests, comp-off claims and ledger, approvals, calendars, availability, and workload warnings", "Advanced HR policy rules, payroll, and attendance integrations"),
        ("Secure client login, project portal, published visibility, UAT", "Client reporting packs and integrations"),
        ("Member, lead, client, and management dashboards", "Forecasting and portfolio scenario planning"),
        ("In-app notifications plus essential email", "Multi-channel rules and integrations"),
    ], [3.25, 3.25])
    d.add_paragraph("Leave and comp-off are confirmed Phase 2 MVP requirements. Payroll, attendance capture, salary processing, and broad HRIS capabilities remain out of scope.")
    d.add_heading("13. Acceptance criteria", 1)
    add_bullets(d, ["A user can create or join multiple teams and hold different roles by scope.", "Any verified team member can submit an organization request, but no tenant resources are provisioned before BQ Labs approval.", "Approval atomically creates the organization and makes the requester Organization Owner and initial Organization Administrator.", "The platform-admin dashboard provides approval, usage, quota, subscription, security, tenant-state, job-health, and immutable admin-action views.", "Every new organization receives the approved starter quotas; warnings appear at 80% and 95%, and hard-limit checks are atomic.", "Reaching a quota blocks only consumption-increasing actions and never deletes existing organization data.", "An expired or payment-suspended organization can read authorized data and contact BQ Labs, but every operational write is rejected server-side.", "An audited BQ Labs administrator action can extend or restore the correct plan capabilities without inconsistent state.", "Phase 2 contains no checkout, payment gateway, invoice, refund, or self-service renewal processing.", "A Security Locked organization cannot regain ordinary access merely through a plan extension.", "Organization Owner and authorized Organization Administrators can appoint administrators and transfer Team Ownership.", "The current Team Owner can transfer ownership of their team, and the system never permits a team or organization to lose its required owner.", "The team creator becomes Team Owner and can assign, replace, or remove the Team Lead.", "Only the Team Owner and currently assigned Team Lead can change client visibility for projects within their authorized scope.", "A removed or transferred Team Lead immediately loses lead-only authority for the former scope, including visibility-policy management.", "A member can submit leave and comp-off evidence; each decision has a reason and immutable audit record.", "For leave starting today or later, Team Lead approval is Level 1 and client approval is Level 2 for every affected project configured to require it.", "The client cannot act before Level 1 approval and sees only the affected project's schedule impact and safe handover information.", "When multiple affected projects require client approval, every required project step must approve before final approval.", "A leave starting before the organization's current local date is system-approved automatically after all validations and is labelled Backdated in the audit trail.", "Backdated auto-approval cannot bypass balance, overlap, authorization, or data-integrity rules.", "Approved leave updates availability and warns about affected assignments without silently reassigning tasks.", "Comp-off utilization cannot exceed approved available balance, and balance is derived from immutable ledger entries.", "Private leave reasons, balances, medical details, comp-off evidence, and unrelated projects are never exposed to clients.", "A client can securely sign in and see tasks only within projects explicitly assigned to that client.", "Each project supports Summarized and Full Detail visibility profiles, with Summarized applied by default.", "Changing a project's visibility profile immediately changes authorized API responses, is audit-logged, and never bypasses mandatory exclusions.", "A client cannot discover a task from another project through search, filters, guessed identifiers, exports, notifications, or direct URLs.", "Cross-tenant and unauthorized object access tests fail closed for every endpoint.", "Invitations expire, are revocable, cannot be replayed, and cannot be redirected to a different account without authorization.", "Every task transition, approval, access change, and client-visible publication is attributable and timestamped.", "Blocked, overdue, unassigned, due-soon, and decision-needed work is discoverable from the relevant dashboards.", "Critical member and client workflows are usable at 360px width and with keyboard-only navigation.", "Collection APIs are filtered, paginated, payload-bounded, and covered by query-count tests.", "Backup restoration, migration rollback, and security release checks are exercised before production launch."])
    d.add_heading("14. Implementation roadmap", 1)
    add_table(d, ["Stage", "Deliverable", "Exit gate"], [
        ("0. Discovery", "Workflow decisions, permission matrix, data dictionary, wireflows, threat model.", "Product decisions and acceptance criteria approved."),
        ("1. Foundation", "Platform admin, organization approval, tenancy, auth, teams, invitations, roles, audit framework, and CI security checks.", "Approval, isolation, and auth negative tests pass."),
        ("2. Delivery core", "Projects, modules, milestones, tasks, assignment, state transitions, comments.", "End-to-end internal workflow passes."),
        ("3. Risk and workforce", "Blockers, change requests, leave, comp-off ledger, approvals, availability, dashboards, and notifications.", "Delivery, leave, balance, approval, and performance tests pass."),
        ("4. Client portal", "Client access, publishing controls, UAT, client dashboard, essential email.", "Client isolation and visibility tests pass."),
        ("5. Subscription readiness", "Plan catalogue, entitlement snapshot, admin-controlled expiry, read-only state, and platform-admin controls; no payment processing.", "Entitlement and restricted-write tests pass."),
        ("6. Hardening", "Accessibility, load tests, backup restore, monitoring, runbooks, pilot feedback.", "Release checklist accepted."),
    ], [.9, 3.7, 1.9])
    d.add_heading("15. Decision register", 1)
    add_table(d, ["Decision", "Status"], [
        ("Clients can authenticate in the Phase 2 MVP.", "Approved"),
        ("Clients can see tasks only for projects explicitly assigned to them.", "Approved"),
        ("Client detail is configurable per project using Summarized or Full Detail visibility profiles.", "Approved"),
        ("The Team Owner and currently assigned Team Lead may manage project visibility within their scope.", "Approved"),
        ("Team Lead is transferable; former lead authority is revoked when the assignment ends.", "Approved"),
        ("The verified organization creator becomes Organization Owner and initial Organization Administrator.", "Approved"),
        ("Organization Owner and authorized Organization Administrators may appoint administrators and transfer Team Ownership.", "Approved"),
        ("Any verified team member may request an organization; BQ Labs approval is required before provisioning.", "Approved"),
        ("A dedicated BQ Labs platform-admin page and dashboard manages approvals and platform-wide administrative actions.", "Approved"),
        ("Starter quotas: 25 users, 5 teams, 10 active projects, 10,000 tasks, and 2 GB attachments per organization.", "Approved"),
        ("Subscription expiry is read-only; Phase 2 restoration is controlled manually by BQ Labs; security lock remains separate.", "Approved"),
        ("Launch packaging uses Pilot, Team, Business, and Enterprise plans with client accounts included.", "Recommended for validation"),
        ("Payment gateway, invoicing, refunds, tax handling, and self-service renewal are deferred to the next phase.", "Approved"),
        ("Leave and earned comp-off workflows are included in the Phase 2 MVP.", "Approved"),
        ("Leave approval is sequential: Team Lead Level 1 and optional Client Level 2 configured per project.", "Approved"),
        ("All required affected clients must approve future/same-day leave; backdated leave is system-auto-approved after validation.", "Approved"),
        ("A client decline sets Changes Required; revisions preserve history and restart the full approval chain.", "Approved"),
        ("Unanswered future/same-day leave stays Pending with reminders and escalation; silence never auto-approves.", "Approved"),
        ("Default leave reminders occur at 24 and 48 hours, with escalation at 72 hours; organization admins may configure them.", "Approved"),
        ("Lead delays escalate to Team Owner; client delays escalate to Team Lead, Team Owner, and secondary client approver.", "Approved"),
        ("Clients may approve only before leave begins; pending required steps become Unapproved Leave at the start deadline.", "Approved"),
        ("Unapproved Leave deducts the requested balance unless cancelled before start; later reversals require audited adjustment.", "Approved"),
        ("Insufficient normal-leave balance does not block recording; the shortfall is reported for transparency. Comp-off cannot exceed earned balance.", "Approved"),
        ("Organization members see name, dates, duration, and status on the shared calendar; private leave data remains restricted.", "Approved"),
        ("Task assignment requires explicit team-member action; proposed assignments remain Pending Member Action until resolved.", "Approved"),
        ("Task-assignment response is due by the next working day calculated from the task's project calendar.", "Approved"),
        ("Missed assignment-response deadlines remain Pending Member Action - Overdue and escalate to Team Lead; no auto-acceptance.", "Approved"),
        ("Team Lead may withdraw or reassign overdue proposals with a mandatory reason and retained history.", "Approved"),
        ("Clients see Not Started until member acceptance; all internal assignment-proposal states remain hidden.", "Approved"),
        ("After acceptance, Summarized shows responsible team; Full Detail may show accepted assignee name.", "Approved"),
        ("Acceptance confirms ownership but task remains Not Started until assignee explicitly selects Start Work.", "Approved"),
        ("Missing planned start sends member/lead alerts but leaves task status unchanged as Not Started.", "Approved"),
        ("Missed-start reminders repeat once per project working day until start, replan, reassignment, or cancellation.", "Approved"),
        ("In-progress updates are optional and event-driven; no mandatory daily progress submission.", "Approved"),
        ("Assigned team members may update all valid task statuses; completed reviews record Reviewed By and review details.", "Approved"),
        ("Any authorized project participant may review a task, including the assignee; reviewer identity and role are recorded.", "Approved"),
        ("Review requirement is task-configurable; required tasks need a current completed review before UAT.", "Approved"),
        ("Task creator becomes Task Owner; only Task Owner may change the task's Review Required setting.", "Approved"),
        ("Organization Admin marking a Team Lead resigned must select a replacement; active task ownership and management duties transfer atomically.", "Approved"),
        ("Departing lead's personally assigned active tasks are proposed to the replacement as Pending Member Action.", "Approved"),
        ("BQ Labs, Quantum OPS, and future products use a reusable Apple-inspired but distinctly BQ Labs-owned design system.", "Approved direction; detailed design workshop pending"),
        ("UAT observations may be raised only by the assigned member or the task's explicitly assigned Tester.", "Approved"),
        ("Working days and leave register are project-scoped; bulk leave data may be uploaded through validated Excel import.", "Approved"),
        ("Summarized is the default; Full Detail remains subject to mandatory security and privacy exclusions.", "Approved"),
        ("MVP deployment model and hosting topology.", "Pending technical review"),
    ], [5.0, 1.5])
    d.add_heading("16. Delivery controls", 1)
    add_bullets(d, ["Do not rewrite Git history.", "Request explicit approval immediately before pushing code to GitHub.", "Request separate explicit approval immediately before deploying to the server.", "Do not commit secrets, credentials, API keys, or .env files."])
    d.add_paragraph("Living-reference control: every approved or modified plan decision must be reflected in both 003.BQLabs-AI-generated.docx and 004.QuantumOPS.docx. Keep the two documents synchronized even when one update is only a governance note or cross-reference.")
    add_consolidated_phase2_decisions(d)
    add_final_design_blueprint(d, compact=False)
    add_secure_self_service_and_admin_update(d)
    d.save(ROOT / "004.QuantumOPS.docx")


if __name__ == "__main__":
    build_bqlabs(); build_quantumops()
